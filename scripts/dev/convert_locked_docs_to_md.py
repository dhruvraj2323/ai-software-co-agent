"""Convert locked DOCX documents to GitHub-readable Markdown mirrors."""

from __future__ import annotations

from pathlib import Path
import re

from docx import Document


ROOT = Path(__file__).resolve().parents[2]
DOCS_DIR = ROOT / "docs"


def clean_text(text: str) -> str:
    """Normalize whitespace while preserving readable text."""
    return re.sub(r"[ \t]+", " ", text).strip()


def heading_level(style_name: str) -> int | None:
    """Return Markdown heading level for common Word heading styles."""
    match = re.match(r"Heading ([1-6])$", style_name or "")
    if match:
        return int(match.group(1))
    return None


def is_list_paragraph(paragraph) -> bool:
    """Detect Word list paragraphs using their style name."""
    style = paragraph.style.name if paragraph.style else ""
    return "List" in style


def convert_table(table) -> str:
    """Convert a Word table to a Markdown table."""
    rows: list[list[str]] = []

    for row in table.rows:
        cells = [clean_text(cell.text).replace("|", r"\|") for cell in row.cells]
        rows.append(cells)

    if not rows:
        return ""

    column_count = max(len(row) for row in rows)
    rows = [row + [""] * (column_count - len(row)) for row in rows]

    lines = [
        "| " + " | ".join(rows[0]) + " |",
        "| " + " | ".join("---" for _ in range(column_count)) + " |",
    ]

    for row in rows[1:]:
        lines.append("| " + " | ".join(row) + " |")

    return "\n".join(lines)


def convert_document(source: Path, destination: Path) -> None:
    """Convert one DOCX document to Markdown."""
    document = Document(source)
    output: list[str] = []

    for block in document.element.body:
        if block.tag.endswith("}p"):
            paragraph = next(
                (p for p in document.paragraphs if p._p is block),
                None,
            )
            if paragraph is None:
                continue

            text = clean_text(paragraph.text)
            if not text:
                continue

            level = heading_level(paragraph.style.name if paragraph.style else "")
            if level is not None:
                output.append(f"{'#' * level} {text}")
            elif is_list_paragraph(paragraph):
                output.append(f"- {text}")
            else:
                output.append(text)

            output.append("")

        elif block.tag.endswith("}tbl"):
            table = next(
                (t for t in document.tables if t._tbl is block),
                None,
            )
            if table is None:
                continue

            table_markdown = convert_table(table)
            if table_markdown:
                output.append(table_markdown)
                output.append("")

    destination.write_text(
        "\n".join(output).rstrip() + "\n",
        encoding="utf-8",
    )


def main() -> None:
    """Convert every locked DOCX under docs/ to a Markdown mirror."""
    documents = sorted(
        path
        for path in DOCS_DIR.rglob("*.docx")
        if "FINAL_LOCKED" in path.name
    )

    if len(documents) != 19:
        raise RuntimeError(
            f"Expected 19 FINAL_LOCKED DOCX files, found {len(documents)}."
        )

    for source in documents:
        destination = source.with_suffix(".md")
        convert_document(source, destination)
        print(f"CONVERTED: {source.relative_to(ROOT)}")
        print(f"       -> {destination.relative_to(ROOT)}")

    print()
    print(f"CONVERSION COMPLETE: {len(documents)} documents")


if __name__ == "__main__":
    main()